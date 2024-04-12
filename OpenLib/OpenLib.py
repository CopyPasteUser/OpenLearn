import json
import os
import threading
import tkinter as tk
from tkinter import messagebox
from pathlib import Path
from docx import Document
from g4f.client import Client  # type: ignore
import math
from PIL import Image, ImageDraw, ImageTk
import requests
import tkinter as tk
from tkinter import messagebox
from urllib.parse import unquote


def draw_english_flag(canvas):
    # Draw white background with thin black border
    canvas.create_rectangle(0, 0, 24, 24, fill="white", outline="black")

    # Draw red cross with thinner lines and thin black border
    canvas.create_rectangle(8, 0, 16, 24, fill="red", outline="black")
    canvas.create_rectangle(0, 8, 24, 16, fill="red", outline="black")

def draw_flag_vertical(canvas, x, y, colors):
    stripe_width = 24 // len(colors)
    for i, color in enumerate(colors):
        # Draw colored stripe with thin black border
        canvas.create_rectangle(x + i * stripe_width, y, x + (i + 1) * stripe_width, y + 24, fill=color, outline="black")

def draw_flag_horizontal(canvas, x, y, colors):
    stripe_height = 24 // len(colors)
    for i, color in enumerate(colors):
        # Draw colored stripe with thin black border
        canvas.create_rectangle(x, y + i * stripe_height, x + 24, y + (i + 1) * stripe_height, fill=color, outline="black")


def sanitize_filename(filename):
    valid_chars = '-.() abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789'
    return ''.join(c if c in valid_chars else '_' for c in filename)

def create_folder(parent_folder, folder_name):
    folder_path = Path(parent_folder) / folder_name
    folder_path.mkdir(parents=True, exist_ok=True)
    return folder_path

def create_file(parent_folder, file_name, content=""):
    parent_path = Path(parent_folder)
    file_path = parent_path / file_name
    with open(str(file_path) + ".txt", 'w', encoding="utf-8") as file:
        file.write(content)
    return file_path

def askAI(message):
    client = Client()
    response = client.chat.completions.create(
        model="openchat_3.5",
        messages=[{"role": "user", "content": message}],
    )
    return response.choices[0].message.content

def checkChapterAmount(experience):  
    if experience == "Elementaryschool (grade 1-4)":
        return 2
    if experience == "Middleschool(grade 5-8)":
        return 3
    if experience == "Highschool(grade 9-12)":
        return 5
    if experience == "University":
        return 7

def getBookOutline(topic, experience,language):
    format = "{ "'BOOKTITLE'": "", "'CHAPTERS'": [ { "'CHAPTERTITLE'": "", "'SUBCHAPTERS'": [ "", "", "", "" ] },"
    chapterAmount = checkChapterAmount(experience)
    prompt = (f"You are a teacher and you are writing a book for your students of {experience} with maximum amount  of {chapterAmount} chapters" +
              "The only response you shall give me is a outline" +
              f"in json format." +
              "it shall only consist of the json part, make it json and not a string!" +
              f"the json keys should be in this exact format: {format} " +
              "(keys must be in all caps, and subchapters must be numerated)  ." + 
              "The book is")
    
    
    
    print(experience)
    print(str(chapterAmount) + " Chapters")
    
    bookOutline = askAI(prompt + f" about {topic}. remember to write it in the {language} language and make sure its grammatically correct!")
    if "json" in bookOutline[:7]:
        bookOutline = bookOutline[7:-3]
    with open("util/bookOutline.json", "w", encoding="utf-8") as json_file:
       json_file.write(bookOutline)

def create_subchapter(doc, subchapter, experience,language,context):
    content = askAI(f"Write the content for the subchapter: {subchapter}. the written langauge is {language}, look out for appropriate language/complexity."+
                    f"when writing this subchapter, remember its mainly about: {context}! also keep in mind that its written for students"+
                    f"in {experience}. only use true information and use correct grammar.")

    # Split content into lines
    lines = content.split("\n")

    # Use the first line as subchapter title
    subchapter_title = lines[0]

    # Add subchapter title as a heading
    doc.add_heading(subchapter_title, level=3)

    # Add remaining lines as content paragraphs
    for line in lines[1:]:
         doc.add_paragraph(line)

def createBook(topic, experience, language):
    getBookOutline(topic, experience, language)
    
    # Open the bookOutline.json file using a context manager to ensure it's properly closed
    with open("util/bookOutline.json", 'r', encoding='utf-8') as f:
        data = json.load(f)
        x = 0
        
        bookTitle = data["BOOKTITLE"]
        sanitizedBookTitle = sanitize_filename(bookTitle)

        # Create a new Word document
        doc = Document()

        # Add book title as the main heading
        doc.add_heading(unquote(bookTitle), level=1)

        for chapter in data['CHAPTERS']:
            x += 1
            
            chapterTitle = str(x) + "." + unquote(chapter['CHAPTERTITLE'])

            # Add chapter title as a heading
            doc.add_heading(chapterTitle, level=2)

            threads = []
            for subchapter in chapter["SUBCHAPTERS"]:
                thread = threading.Thread(target=create_subchapter, args=(doc, subchapter, experience, language,sanitizedBookTitle))
                threads.append(thread)
                thread.start()

            # Wait for all threads to finish
            for thread in threads:
                thread.join()

        # Save the Word document
        doc.save(os.path.join(".", sanitizedBookTitle + ".docx"))


def run_gui():
    # Create the main window
    window = tk.Tk()
    window.title("Book Creation")

    # Create a frame for topic entry
    topic_frame = tk.Frame(window)
    topic_frame.pack(pady=10)

    # Topic label and entry
    topic_label = tk.Label(topic_frame, text="Enter the topic:")
    topic_label.grid(row=0, column=0, sticky="w")
    topic_entry = tk.Entry(topic_frame, width=50)
    topic_entry.grid(row=0, column=1)

    # Create a frame for experience level selection
    experience_frame = tk.Frame(window)
    experience_frame.pack(pady=10)

    # Experience level label
    experience_label = tk.Label(experience_frame, text="Select the experience level:")
    experience_label.grid(row=0, column=0, sticky="w")

    # Create a variable to store the selected experience level
    selected_experience = tk.StringVar(window)

    # Experience levels
    experience_levels = ["Elementaryschool (grade 1-4)", "Middleschool(grade 5-8)", "Highschool(grade 9-12)", "University"]

    # Create radio buttons for experience levels
    for i, exp in enumerate(experience_levels):
        rb = tk.Radiobutton(experience_frame, text=exp, variable=selected_experience, value=exp)
        rb.grid(row=i+1, column=0, sticky="w")

    # Create a frame for language selection
    language_frame = tk.Frame(window)
    language_frame.pack(pady=10)

    # Language label
    language_label = tk.Label(language_frame, text="Select the language:")
    language_label.grid(row=0, column=0, sticky="w")

    # Create a variable to store the selected language
    selected_language = tk.StringVar(window)

    # Language buttons with corresponding flag images
    languages = ["english", "german", "french", "spanish", "italian"]  # ISO 3166-1 alpha-2 country codes

    # Flag colors for each language
    flag_colors = {
        "english": ["white", "red"],          # English flag: White with a red cross
        "german": ["black", "red", "yellow"],# German flag: Black, red, yellow horizontal stripes
        "french": ["blue", "white", "red"],  # French flag: Blue, white, red vertical stripes
        "spanish": ["red", "yellow", "red"],  # Spanish flag: Red, yellow, red horizontal stripes
        "italian": ["green", "white", "red"]  # Italian flag: Green, white, red vertical stripes
    }

     # Create language buttons with flags
    for i, lang in enumerate(languages):
        flag_colors_for_lang = flag_colors.get(lang)
        if flag_colors_for_lang:
            lang_btn = tk.Button(language_frame, text=lang.upper(), bg="white", command=lambda l=lang: selected_language.set(l))
            lang_btn.grid(row=0, column=i+1, padx=5, pady=5)
            canvas = tk.Canvas(language_frame, width=24, height=24)
            canvas.grid(row=1, column=i+1)
            if lang == "english":  # Draw English flag
                draw_english_flag(canvas)
            elif lang in ["german", "spanish"]:  
                draw_flag_horizontal(canvas, 0, 0, flag_colors_for_lang)
            else:
                draw_flag_vertical(canvas, 0, 0, flag_colors_for_lang)

    def start_book_creation():
        topic = topic_entry.get()
        experience = selected_experience.get()
        language = selected_language.get()
        if not topic.strip():
            messagebox.showwarning("Warning", "Please enter a topic.")
        elif not experience:
            messagebox.showwarning("Warning", "Please select an experience level.")
        elif not language:
            messagebox.showwarning("Warning", "Please select a language.")
        else:
            # Start book creation process
            createBook(topic, experience, language)
            messagebox.showinfo("Information", f"Book creation for '{topic}' at '{experience}' level in '{language.upper()} language.")

    # Start button
    start_button = tk.Button(window, text="Start Book Creation", command=start_book_creation)
    start_button.pack(pady=10)

    # Run the GUI event loop
    window.mainloop()


def main():
    gui_thread = threading.Thread(target=run_gui)
    gui_thread.start()

if __name__ == "__main__":
    main()
