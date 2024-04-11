import json
import os
import threading
import tkinter as tk
from tkinter import messagebox
from pathlib import Path
from docx import Document
from g4f.client import Client  # type: ignore
import math

def is_supported_image(file_path):
    supported_formats = [".jpg", ".jpeg", ".png", ".gif"]
    return any(file_path.lower().endswith(fmt) for fmt in supported_formats)

def find_cover_picture(directory):
    for filename in os.listdir(directory):
        if not filename.endswith(".json"):
            return f"{filename}"
    return None

def sanitize_filename(filename):
    valid_chars = '-.() abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789'
    return ''.join(c if c in valid_chars else '_' for c in filename)

def create_folder(parent_folder, folder_name):
    folder_path = Path(parent_folder) / folder_name
    folder_path.mkdir(parents=True, exist_ok=True)
    return folder_path

def create_file(parent_folder, file_name, content=""):
    parent_path = Path(parent_folder)
    file_path = parent_path / file_name
    with open(str(file_path) + ".txt", 'w', encoding="utf-8") as file:
        file.write(content)
    return file_path

def askAI(message):
    client = Client()
    response = client.chat.completions.create(
        model="openchat_3.5",
        messages=[{"role": "user", "content": message}],
    )
    return response.choices[0].message.content

def checkChapterAmount(experience):
    if experience == "Elementaryschool":
        return 2
    if experience == "Middleschool":
        return 3
    if experience == "Highschool":
        return 5
    if experience == "University":
        return 7

def getBookOutline(topic, experience):
    prompt = ("You are a teacher and you are writing a book for your pupils." +
              "The only response you shall give me is a outline in Chapters and subchapters" +
              "from the book in json fromat. Please make sure to make the complexity and length adequate for the audience. i.e the younger the audience the simpler and shorter" +
              "it shall only consist of the json part, make it json and not a string!" +
              "the format should be BOOKTITLE -> CHAPTERS -> CHAPTERTITLE -> SUBCHAPTERS 1.1, 1.2 etc" +
              "(they must not have a key for number but must be in all caps)  ." +
              "The book is")
    
    chapterAmount = checkChapterAmount(experience)
    subchapterAmount = math.ceil(chapterAmount/2)
    print(experience)
    print(str(chapterAmount) + " Chapters")
    print(str(subchapterAmount) + " Subchapters")
    bookOutline = askAI(prompt + f" about {topic} for {experience} with maximum amount  of {chapterAmount} of chapters and {subchapterAmount} of subchapters.")
    if "json" in bookOutline[:7]:
        bookOutline = bookOutline[7:-3]
    with open("util/bookOutline.json", "w", encoding="utf-8") as json_file:
       json_file.write(bookOutline)

def create_subchapter(doc, subchapter, experience):
    content = askAI(f"Write the content for the subchapter: {subchapter} with appropriate language/complexity for {experience}. And also make sure only to include true information and explain well. and dont split this subchapter into more subchapters!!! and make sure to remember in which context you are writing this")

    # Split content into lines
    lines = content.split("\n")

    # Use the first line as subchapter title
    subchapter_title = lines[0]

    # Add subchapter title as a heading
    doc.add_heading(subchapter_title, level=3)

    # Add remaining lines as content paragraphs
    for line in lines[1:]:
         doc.add_paragraph(line)

def createBook(topic, experience):
    getBookOutline(topic, experience)
    
    # Open the bookOutline.json file using a context manager to ensure it's properly closed
    with open("util/bookOutline.json", 'r') as f:
        data = json.load(f)
        x = 0
        
        bookTitle = data["BOOKTITLE"]
        sanitizedBookTitle = sanitize_filename(bookTitle)
        
        # Create a cover picture path
        cover_picture_path = find_cover_picture("util")
        
        # Create a new Word document
        doc = Document()
        
        print(cover_picture_path)
        # Add cover picture if found
        if cover_picture_path and is_supported_image(cover_picture_path):
            doc.add_picture("util/" + cover_picture_path)  # Adjust width as needed
        else:
            print("Warning: Cover picture not found or unsupported format.")
        
        # Add book title as the main heading
        doc.add_heading(bookTitle, level=1)
        print(bookTitle)
        
        for chapter in data['CHAPTERS']:
            x += 1
            chapterTitle = str(x) + "." + chapter['CHAPTERTITLE']
            
            # Add chapter title as a heading
            doc.add_heading(chapterTitle, level=2)
            print(chapterTitle)
            
            threads = []
            for subchapter in chapter["SUBCHAPTERS"]:
                thread = threading.Thread(target=create_subchapter, args=(doc, subchapter, experience))
                threads.append(thread)
                thread.start()
            
            # Wait for all threads to finish
            for thread in threads:
                thread.join()

        # Save the Word document
        doc.save(os.path.join(".", sanitizedBookTitle + ".docx"))  # Save in the current directory

def run_gui():
    # Create the main window
    window = tk.Tk()
    window.title("Book Creation")

    # Create a frame for topic entry
    topic_frame = tk.Frame(window)
    topic_frame.pack(pady=10)

    # Topic label and entry
    topic_label = tk.Label(topic_frame, text="Enter the topic:")
    topic_label.grid(row=0, column=0, sticky="w")
    topic_entry = tk.Entry(topic_frame, width=50)
    topic_entry.grid(row=0, column=1)

    # Create a frame for experience level selection
    experience_frame = tk.Frame(window)
    experience_frame.pack(pady=10)

    # Experience level label
    experience_label = tk.Label(experience_frame, text="Select the experience level:")
    experience_label.grid(row=0, column=0, sticky="w")

    # Create a variable to store the selected experience level
    experience_var = tk.StringVar(window)

    # Experience levels
    experience_levels = ["Elementaryschool", "Middleschool", "Highschool", "University"]

    # Create radio buttons for experience levels
    for i, exp in enumerate(experience_levels):
        rb = tk.Radiobutton(experience_frame, text=exp, variable=experience_var, value=exp)
        rb.grid(row=i+1, column=0, sticky="w")

    def start_book_creation():
        topic = topic_entry.get()
        selected_experience = experience_var.get()
        if topic.strip() == "":
            messagebox.showwarning("Warning", "Please enter a topic.")
        elif selected_experience == "":
            messagebox.showwarning("Warning", "Please select an experience level.")
        else:
            # Start book creation process
            createBook(topic, selected_experience)
            messagebox.showinfo("Information", "Book creation completed successfully.")
            window.destroy()

    # Start button
    start_button = tk.Button(window, text="Start Book Creation", command=start_book_creation)
    start_button.pack(pady=10)

    # Run the GUI event loop
    window.mainloop()

def main():
    gui_thread = threading.Thread(target=run_gui)
    gui_thread.start()

if __name__ == "__main__":
    main()
